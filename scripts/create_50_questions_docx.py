# scripts/create_50_questions_docx.py
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CG e-Procurement Chatbot")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Dark Blue
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("50 Scenario-Based Test Questions (Store & Purchase Rules)")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Gray
    
    doc.add_paragraph() # Spacer
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "This document contains a set of 50 realistic, scenario-based test questions designed to evaluate "
        "and test the CG e-Procurement Chatbot. The questions cover Store Purchase Rules of Chhattisgarh, "
        "General Financial Rules (GFR), CVC guidelines, bid evaluation parameters, MSME exemptions, and e-procurement portal "
        "operations. Questions are provided in a mix of English, Hindi, and Hinglish (transliterated Hindi) to test multilingual capability."
    )
    intro_run.font.size = Pt(11)
    
    doc.add_paragraph() # Spacer
    
    # Data structure
    questions_data = {
        "A. Store Purchase Rules & Procurement Planning": [
            ("Office ko ₹50,000 ka printer cartridge purchase karna hai. Chhattisgarh Store Purchase Rules ke under direct purchase allowed hai kya?", "Hinglish"),
            ("What is the maximum duration for which a short-term tender notice of 2 days can be issued under CG procurement regulations?", "English"),
            ("क्या छत्तीसगढ़ भंडार क्रय नियमों के अंतर्गत स्थानीय उद्योगों (Local MSEs) को मूल्य वरीयता (Price Preference) का लाभ मिलता है?", "Hindi"),
            ("Under what circumstances can a department issue a limited tender without publishing a Notice Inviting Tender (NIT) in local newspapers?", "English"),
            ("Hume ₹15 lakh ka networking project initiate karna hai, isme Limited Tender works rule apply hoga ya Goods procurement rule?", "Hinglish"),
            ("क्या भंडार क्रय नियमों के अनुसार तकनीकी रूप से अयोग्य निविदाकर्ता की वित्तीय निविदा खोली जा सकती है?", "Hindi"),
            ("Can a department purchase items from a cooperative society or public sector undertaking directly, and is there any value limit under CG rules?", "English"),
            ("Humare office ko aane wale financial year ke liye annual maintenance services plan karni hai. Estimated budget kaise calculate karein?", "Hinglish"),
            ("क्या आपातकालीन चिकित्सा उपकरणों की खरीद के लिए भी भंडार क्रय नियम (Store Purchase Rules) के तहत खुली निविदा आवश्यक है?", "Hindi"),
            ("Is there a specific threshold value in Chhattisgarh Store Purchase Rules above which procurement through the state e-procurement portal becomes mandatory?", "English"),
        ],
        "B. GFR, Financial Sanctions & CVC Vigilance Guidelines": [
            ("GFR Rule 144(xi) restricts bidders from countries sharing land borders with India. Does this apply to state-level sub-contracts as well?", "English"),
            ("Humare project ka administrative approval toh mil gaya hai, par actual purchase se pehle final financial sanction kis authority se leni chahiye?", "Hinglish"),
            ("क्या निविदा खुलने के बाद बोली लगाने वालों के बीच दरें (rates) बदलने के लिए कार्टेल (Cartelization) की जांच सीवीसी द्वारा की जाती है?", "Hindi"),
            ("Under what GFR rule can a Proprietary Article Certificate (PAC) be signed, and what is its validity period?", "English"),
            ("Agar emergency me COVID test kits kharidne padein bina competitive bidding ke, to kis GFR rule ke under justify karna hoga?", "Hinglish"),
            ("क्या बिना प्रशासनिक स्वीकृति (Administrative Approval) के निविदा प्रक्रिया (tender process) प्रारंभ की जा सकती है?", "Hindi"),
            ("What are the CVC guidelines on the validity and extension of bank guarantees submitted as performance security?", "English"),
            ("Single tender invite karne par agar sirf wahi purana vendor response kare, to kya uski rate direct accept kar sakte hain?", "Hinglish"),
            ("सीवीसी (CVC) निर्देशों के अनुसार निविदा में 'Negotiations' (वार्ता) केवल न्यूनतम बोलीदाता (L1) के साथ ही क्यों की जा सकती है?", "Hindi"),
            ("Does GFR allow purchase of goods through a search on GeM using the "
             "\"Custom Bid\" feature when standard specifications do not match?", "English"),
        ],
        "C. Specifications, Bidder Eligibility & MSME Exemptions": [
            ("Specs me \"high-quality RAM\" ya \"fast processor\" jaise ambiguous terms use karne par CVC guidelines kya kehti hain?", "Hinglish"),
            ("क्या किसी निविदा में निविदाकर्ताओं के लिए न्यूनतम वार्षिक टर्नओवर (Minimum Annual Turnover) का मानदंड रखना अनिवार्य है?", "Hindi"),
            ("Can we require the bidder to have a local service center in Chhattisgarh to qualify for the technical evaluation?", "English"),
            ("MSME certificate hone par EMD exemption to mil jata hai, par kya performance security (SD) se bhi exemption milta hai?", "Hinglish"),
            ("क्या निविदा शुरू होने के बाद किसी एक निविदाकर्ता के अनुरोध पर अर्हता शर्तों (eligibility criteria) को शिथिल किया जा सकता है?", "Hindi"),
            ("How should we handle a bid where the MSME certificate is valid but does not cover the specific category of items being procured?", "English"),
            ("Technical bid sheet me bid capability verify karne ke liye kya past performance report accept karni chahiye?", "Hinglish"),
            ("यदि बोलीदाता ने मूल दस्तावेज अपलोड नहीं किए हैं, केवल स्व-सत्यापित (self-attested) प्रतियां दी हैं, तो क्या बोली स्वीकार्य है?", "Hindi"),
            ("Under GFR rules, are startup entities registered with DPIIT exempted from the \"prior experience\" criteria if they meet quality specifications?", "English"),
            ("Department laptop specification me screen resolution exact 1920x1080 px likh sakta hai ya ranges dena zaroori hai?", "Hinglish"),
        ],
        "D. Bid Evaluation, Negotiations (L1) & Contract Award": [
            ("If the L1 bidder backs out after the financial evaluation, can we award the contract to the L2 bidder at L1 rates?", "English"),
            ("Technical comparative statement verify karte waqt agar kisi point par contradiction ho, to evaluation committee kya step legi?", "Hinglish"),
            ("क्या बोलीदाता को तकनीकी मूल्यांकन (Technical Evaluation) में असफल होने का कारण बताना विभाग के लिए अनिवार्य है?", "Hindi"),
            ("What is the process to be followed if the lowest bid (L1) is received from a joint venture or consortium instead of an individual firm?", "English"),
            ("Final award se pehle agar department ko pata chale ki L1 bidder ne fake experience certificate submit kiya hai, to kya action lena hoga?", "Hinglish"),
            ("क्या वित्तीय बोलियां (Financial Bids) खोलने के तुरंत बाद कार्य आदेश (Work Order) जारी किया जा सकता है?", "Hindi"),
            ("Under what circumstances can a purchase order be amended after it has been signed by both parties?", "English"),
            ("Agar multiple bidders ke quotes exact same ho (Tied L1), to tender award karne ke liye tie-breaker rule kya hai?", "Hinglish"),
            ("क्या बोलीदाता द्वारा जमा की गई बैंक गारंटी (Performance Security) को अनुबंध के उल्लंघन पर जब्त (forfeit) किया जा सकता है?", "Hindi"),
            ("What is the maximum percentage of delivery delay penalty (liquidated damages) that can be levied on a defaulting vendor under standard contracts?", "English"),
        ],
        "E. Portal Operations, DSC & Technical Troubleshooting": [
            ("New vendor registration ke time portal par DSC registration verify karte waqt 'Java error' aaye to browser settings me kya changes karein?", "Hinglish"),
            ("क्या बोलीदाता निविदा जमा करने की अंतिम तिथि समाप्त होने के बाद भी अपने दस्तावेज़ संशोधित कर सकता है?", "Hindi"),
            ("How long does the portal take to automatically initiate the online EMD refund for bidders who are rejected at the technical round?", "English"),
            ("Portal par bid decrypt karne ke liye technical opener ko personal Class-III DSC check karna kyun mandatory hota hai?", "Hinglish"),
            ("क्या शुद्धिपत्र (Corrigendum) जारी करने पर निविदा जमा करने की अंतिम तिथि को बढ़ाना अनिवार्य है?", "Hindi"),
            ("If the online payment gateway fails during the payment of tender document fees, how can the bidder retrieve the transaction receipt?", "English"),
            ("Edge browser setup guide ke according dynamic links open karne ke liye 'IE Mode' active hona kyun required hai?", "Hinglish"),
            ("क्या टेंडर ऑपरेटर तकनीकी बोली खोलने (Technical Bid Opening) के लिए स्वयं उत्तरदायी है या इसमें अप्रूवर की सहमति भी चाहिए?", "Hindi"),
            ("What is the procedure on the CHiPS e-procurement portal if a bidder wants to submit EMD via NEFT/RTGS challan instead of net banking?", "English"),
            ("Portal par commercial schedule (BOQ.xls) fill karte waqt formula modification error aaye to vendor ko kya karna chahiye?", "Hinglish"),
        ]
    }
    
    global_q_index = 1
    
    for category, list_of_q in questions_data.items():
        # Category Heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(category)
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)
        
        # Table of Questions
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'S.No.'
        hdr_cells[1].text = 'Question Scenario'
        hdr_cells[2].text = 'Language'
        
        # Format Headers
        for cell in hdr_cells:
            set_cell_background(cell, "1F4E79")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
        # Set Widths
        widths = [Inches(0.8), Inches(4.7), Inches(1.0)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                
        # Populate Questions
        for q_text, q_lang in list_of_q:
            row_cells = table.add_row().cells
            row_cells[0].text = str(global_q_index)
            row_cells[1].text = q_text
            row_cells[2].text = q_lang
            
            # Format row font/alignment
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set Widths for new row cells
            for idx, width in enumerate(widths):
                row_cells[idx].width = width
                
            global_q_index += 1
            
        doc.add_paragraph() # Spacer between categories
        
    output_path = "c:/cg-eproc-chatbot/scripts/test_questions_50.docx"
    doc.save(output_path)
    print(f"Document successfully created and saved to {output_path}")

if __name__ == "__main__":
    create_document()
