# scripts/convert_evaluation_to_docx.py
import os
import json
import re
from datetime import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_paragraph_with_spacing(doc, text="", style='Normal', space_before=0, space_after=6, line_spacing=1.15):
    """Helper to add paragraph with fine control over spacing."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        p.add_run(text)
    return p

def format_bold_runs(paragraph, text):
    """Parse text for markdown-style **bolding** and apply as runs in docx."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def create_docx_report():
    json_path = os.path.join("scripts", "dual_evaluation_report.json")
    docx_path = os.path.join("scripts", "dual_evaluation_report.docx")
    
    if not os.path.exists(json_path):
        print(f"Error: {json_path} not found.")
        return
        
    with open(json_path, "r", encoding="utf-8") as f:
        results = json.load(f)
        
    local_lats = [r["local_latency_ms"] for r in results]
    sarvam_lats = [r["sarvam_latency_ms"] for r in results]
    
    avg_local = int(sum(local_lats) / len(local_lats)) if local_lats else 0
    p95_local = sorted(local_lats)[int(len(local_lats) * 0.95)] if local_lats else 0
    
    avg_sarvam = int(sum(sarvam_lats) / len(sarvam_lats)) if sarvam_lats else 0
    p95_sarvam = sorted(sarvam_lats)[int(len(sarvam_lats) * 0.95)] if sarvam_lats else 0
    
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set Normal Font Style
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("CG e-Procurement Chatbot")
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Dark Blue
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Dual Model Performance & Response Evaluation Report")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x70, 0x30, 0xA0) # Deep Purple
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(24)
    run_meta = p_meta.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Test Set: 50 Questions")
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    
    # --- Performance Summary Table ---
    h_summary = doc.add_paragraph()
    h_summary.paragraph_format.space_before = Pt(12)
    h_summary.paragraph_format.space_after = Pt(6)
    run_h_summary = h_summary.add_run("Model Configuration & Latency Summary")
    run_h_summary.font.bold = True
    run_h_summary.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run_h_summary.font.size = Pt(14)
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Shading Accent 1'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model / Parameter'
    hdr_cells[1].text = 'Avg Latency (ms)'
    hdr_cells[2].text = 'P95 Latency (ms)'
    hdr_cells[3].text = 'Model Type'
    
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                
    # Data Rows
    row1 = table.rows[1].cells
    row1[0].text = 'Local Model (Ollama)'
    row1[1].text = f"{avg_local} ms"
    row1[2].text = f"{p95_local} ms"
    row1[3].text = 'cg-procurement-chatbot (RAG + Local Model)'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Sarvam 30B Model'
    row2[1].text = f"{avg_sarvam} ms"
    row2[2].text = f"{p95_sarvam} ms"
    row2[3].text = 'sarvam-30b API (Direct Prompting)'
    
    # Stylize data rows
    for r_idx, r in enumerate(table.rows[1:], 1):
        bg_color = "F2F4F7" if r_idx % 2 == 1 else "FFFFFF"
        for cell in r.cells:
            set_cell_background(cell, bg_color)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    
    doc.add_paragraph() # Spacer
    doc.add_page_break()
    
    # --- Questions Section ---
    sections = {
        "A": "Section A: Store Purchase Rules & Procurement Planning",
        "B": "Section B: GFR, Approvals & Financial Control",
        "C": "Section C: Specifications, Bidder Eligibility & MSME Exemptions",
        "D": "Section D: Bid Evaluation, Negotiations (L1) & Contract Award",
        "E": "Section E: Portal Operations, DSC & Technical Troubleshooting"
    }
    
    current_sec = None
    for r in results:
        if r["section"] != current_sec:
            current_sec = r["section"]
            sec_title = sections.get(current_sec, current_sec)
            h_sec = doc.add_paragraph()
            h_sec.paragraph_format.space_before = Pt(24)
            h_sec.paragraph_format.space_after = Pt(12)
            run_h_sec = h_sec.add_run(sec_title)
            run_h_sec.font.size = Pt(16)
            run_h_sec.font.bold = True
            run_h_sec.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            
            # Colored horizontal line below section heading
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_after = Pt(12)
            p_line.paragraph_format.space_before = Pt(0)
            run_line = p_line.add_run("―" * 60)
            run_line.font.bold = True
            run_line.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            
        # Question Header
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(12)
        p_q.paragraph_format.space_after = Pt(3)
        run_q_id = p_q.add_run(f"Question {r['id']}: ")
        run_q_id.font.bold = True
        run_q_id.font.size = Pt(11)
        run_q_id.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
        
        run_q_text = p_q.add_run(r["question"])
        run_q_text.font.bold = True
        run_q_text.font.size = Pt(11)
        
        # Expected Role
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_after = Pt(6)
        p_role.paragraph_format.space_before = Pt(0)
        run_role_lbl = p_role.add_run("Expected Actor Role: ")
        run_role_lbl.font.italic = True
        run_role_lbl.font.size = Pt(9.5)
        run_role_val = p_role.add_run(r["actor"])
        run_role_val.font.bold = True
        run_role_val.font.size = Pt(9.5)
        run_role_val.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        
        # --- Local Answer Block ---
        p_local_title = doc.add_paragraph()
        p_local_title.paragraph_format.space_before = Pt(6)
        p_local_title.paragraph_format.space_after = Pt(2)
        run_local_title = p_local_title.add_run("💻 Local Model (cg-procurement-chatbot)")
        run_local_title.font.bold = True
        run_local_title.font.size = Pt(10)
        run_local_title.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E)
        
        # Local Meta
        p_local_meta = doc.add_paragraph()
        p_local_meta.paragraph_format.space_before = Pt(0)
        p_local_meta.paragraph_format.space_after = Pt(4)
        sources_str = ", ".join(r["local_sources"]) if r["local_sources"] else "None Cited"
        run_local_meta = p_local_meta.add_run(
            f"Latency: {r['local_latency_ms']} ms | Cited Sources: {sources_str}"
        )
        run_local_meta.font.size = Pt(8.5)
        run_local_meta.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        
        # Local Answer Content
        p_local_ans = doc.add_paragraph()
        p_local_ans.paragraph_format.left_indent = Inches(0.2)
        p_local_ans.paragraph_format.space_after = Pt(8)
        p_local_ans.paragraph_format.line_spacing = 1.15
        format_bold_runs(p_local_ans, r["local_answer"])
        for run in p_local_ans.runs:
            run.font.size = Pt(9.5)
            
        # --- Sarvam Answer Block ---
        p_sarvam_title = doc.add_paragraph()
        p_sarvam_title.paragraph_format.space_before = Pt(6)
        p_sarvam_title.paragraph_format.space_after = Pt(2)
        run_sarvam_title = p_sarvam_title.add_run("🇮🇳 Sarvam 30B API")
        run_sarvam_title.font.bold = True
        run_sarvam_title.font.size = Pt(10)
        run_sarvam_title.font.color.rgb = RGBColor(0xC5, 0x5A, 0x11) # Orange Accent
        
        # Sarvam Meta
        p_sarvam_meta = doc.add_paragraph()
        p_sarvam_meta.paragraph_format.space_before = Pt(0)
        p_sarvam_meta.paragraph_format.space_after = Pt(4)
        run_sarvam_meta = p_sarvam_meta.add_run(
            f"Latency: {r['sarvam_latency_ms']} ms | Tokens Generated: {r['sarvam_tokens']}"
        )
        run_sarvam_meta.font.size = Pt(8.5)
        run_sarvam_meta.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        
        # Sarvam Answer Content
        p_sarvam_ans = doc.add_paragraph()
        p_sarvam_ans.paragraph_format.left_indent = Inches(0.2)
        p_sarvam_ans.paragraph_format.space_after = Pt(16)
        p_sarvam_ans.paragraph_format.line_spacing = 1.15
        format_bold_runs(p_sarvam_ans, r["sarvam_answer"])
        for run in p_sarvam_ans.runs:
            run.font.size = Pt(9.5)
            
        # Add visual separator line
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(4)
        p_sep.paragraph_format.space_after = Pt(12)
        run_sep = p_sep.add_run("―" * 50)
        run_sep.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)
        
    doc.save(docx_path)
    print(f"Successfully created and saved Word document comparison report to: {docx_path}")

if __name__ == "__main__":
    create_docx_report()
